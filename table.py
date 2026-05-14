#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
中建四局公文格式化 — 表格处理

包含：
- calc_smart_col_widths(): 智能列宽计算
- extract_table_data(): 从原始 XML 提取表格文本/底色/字体颜色
- rebuild_table(): 在输出文档中重建表格
"""

import re
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from constants import (
    FONT_FANGSONG, SIZE_XIAOSI,
    TABLE_WIDTH_PCT, TABLE_SEQ_COL_WEIGHT, TABLE_COL_WEIGHTS,
    TABLE_WIDTHS_DXA, TABLE_ROW_HEIGHT_TWIPS, TABLE_CELL_SPACING_TWIPS,
)
from detect import clean_text, set_run_font, set_para_spacing


SIZE_XIAOSI_PT = Pt(SIZE_XIAOSI)


def calc_smart_col_widths(rows_data, num_cols):
    """根据每列最大字符数计算列宽权重"""
    col_max_chars = [0] * num_cols
    for row_data in rows_data:
        for j in range(min(len(row_data), num_cols)):
            col_max_chars[j] = max(col_max_chars[j], len(row_data[j].strip()))
    weights = []
    for j in range(num_cols):
        max_len = col_max_chars[j]
        is_seq_col = True
        for row_data in rows_data:
            cell = row_data[j].strip() if j < len(row_data) else ''
            if cell and not re.match(
                r'^(\d{1,3}[\.\-]?\d{0,2}|第[一二三四五六七八九十\d]+项?|序号|编号|No\.?|ID|项)$', cell
            ):
                is_seq_col = False
                break
        if is_seq_col:
            weights.append(TABLE_SEQ_COL_WEIGHT)
        elif max_len <= 6:
            weights.append(TABLE_COL_WEIGHTS[0])
        elif max_len <= 12:
            weights.append(TABLE_COL_WEIGHTS[1])
        elif max_len <= 20:
            weights.append(TABLE_COL_WEIGHTS[2])
        elif max_len <= 35:
            weights.append(TABLE_COL_WEIGHTS[3])
        else:
            weights.append(TABLE_COL_WEIGHTS[4])
    return weights


def extract_table_data(elem):
    """从原始文档的表格 XML 元素提取数据

    返回：('tbl', rows_data, rows_shading, rows_font_color) 或 None
    - rows_data: [[cell_text, ...], ...]
    - rows_shading: [[cell_shading_color, ...], ...]
    - rows_font_color: [[cell_font_color, ...], ...]
    """
    rows_data = []
    rows_shading = []
    rows_font_color = []

    for tr in elem.iter(qn('w:tr')):
        cells = []
        shadings = []
        font_colors = []
        for tc in tr.iter(qn('w:tc')):
            cell_text = ''.join(t.text for t in tc.iter(qn('w:t')) if t.text)
            cells.append(cell_text.strip())

            # 提取原始单元格底色
            tcPr = tc.find(qn('w:tcPr'))
            cell_shading = None
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    cell_shading = shd.get(qn('w:fill'))
            shadings.append(cell_shading)

            # 提取原始单元格字体颜色（取第一个 run 的颜色）
            cell_font_color = None
            for p_elem in tc.iter(qn('w:p')):
                for r_elem in p_elem.iter(qn('w:r')):
                    rPr = r_elem.find(qn('w:rPr'))
                    if rPr is not None:
                        color_elem = rPr.find(qn('w:color'))
                        if color_elem is not None:
                            val = color_elem.get(qn('w:val'))
                            if val:
                                cell_font_color = val
                                break
                if cell_font_color:
                    break
            font_colors.append(cell_font_color)

        if cells:
            rows_data.append(cells)
            rows_shading.append(shadings)
            rows_font_color.append(font_colors)

    if rows_data:
        return ('tbl', rows_data, rows_shading, rows_font_color)
    return None


def rebuild_table(doc, rows_data, rows_shading=None, rows_font_color=None):
    """在输出文档中重建表格

    全局规则：
    - 保持原文表格底稿颜色不变
    - 保持原文表格字体颜色不变
    - 表头行加粗
    - 序号列/短文本居中，纯数字右对齐，其余左对齐
    """
    num_cols = max(len(row) for row in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表格整体宽度
    tbl_elem = table._tbl
    tblPr = tbl_elem.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_elem.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.insert(0, tblW)
    tblW.set(qn('w:w'), str(TABLE_WIDTH_PCT))
    tblW.set(qn('w:type'), 'pct')

    # 列宽
    col_widths = calc_smart_col_widths(rows_data, num_cols)
    total_page_width_dxa = TABLE_WIDTHS_DXA
    total_weight = sum(col_widths)
    for j, weight in enumerate(col_widths):
        width_dxa = int(total_page_width_dxa * weight / total_weight)
        for row in table.rows:
            cell = row.cells[j]
            tc_elem = cell._tc
            tcPr = tc_elem.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc_elem.insert(0, tcPr)
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(width_dxa))
            tcW.set(qn('w:type'), 'dxa')

    # 填充单元格
    for ri, row_data in enumerate(rows_data):
        is_header = (ri == 0)
        tr = table.rows[ri]._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)
        trHeight = trPr.find(qn('w:trHeight'))
        if trHeight is None:
            trHeight = OxmlElement('w:trHeight')
            trPr.append(trHeight)
        trHeight.set(qn('w:val'), str(TABLE_ROW_HEIGHT_TWIPS))
        trHeight.set(qn('w:hRule'), 'atLeast')

        for j, cell_text in enumerate(row_data):
            cell = table.cell(ri, j)
            cell.text = ''
            p = cell.paragraphs[0]
            set_para_spacing(p, twips=TABLE_CELL_SPACING_TWIPS)
            cleaned = clean_text(cell_text)
            run = p.add_run(cleaned)

            # 全局规则4：保持原文表格字体颜色不变
            orig_font_color = (
                rows_font_color[ri][j]
                if rows_font_color and ri < len(rows_font_color) and j < len(rows_font_color[ri])
                else None
            )
            if orig_font_color:
                set_run_font(run, FONT_FANGSONG, SIZE_XIAOSI_PT, bold=is_header, color=RGBColor.from_string(orig_font_color))
            else:
                set_run_font(run, FONT_FANGSONG, SIZE_XIAOSI_PT, bold=is_header)

            # 全局规则2：保持原文表格底稿颜色不变
            orig_shading = (
                rows_shading[ri][j]
                if rows_shading and ri < len(rows_shading) and j < len(rows_shading[ri])
                else None
            )
            if orig_shading:
                tc_elem = cell._tc
                tcPr2 = tc_elem.find(qn('w:tcPr'))
                if tcPr2 is None:
                    tcPr2 = OxmlElement('w:tcPr')
                    tc_elem.insert(0, tcPr2)
                shading = tcPr2.find(qn('w:shd'))
                if shading is None:
                    shading = OxmlElement('w:shd')
                    tcPr2.append(shading)
                shading.set(qn('w:fill'), orig_shading)
                shading.set(qn('w:val'), 'clear')

            # 对齐方式
            header_text = rows_data[0][j].strip() if rows_data and len(rows_data[0]) > j else ''
            is_seq_col = header_text in ('序号', '编号')
            stripped = cleaned.strip()
            if is_header or is_seq_col or len(stripped) <= 20:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif re.match(r'^[\d,.\-+%]+$', stripped) and len(stripped) >= 3:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
