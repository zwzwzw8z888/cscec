#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
中建四局公文格式化 — 批注写入 + 页码

包含：
- _split_run_at(): 在 run 内部指定位置拆分为两个 run
- _add_highlight_to_run(): 给 run 添加高亮底色（当前停用）
- _apply_comments_to_doc(): 统一在文档中添加批注
- _add_page_number(): 添加页码
"""

import re
import datetime
from lxml import etree
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

from constants import FONT_FANGSONG


def _split_run_at(target_elem, run, split_pos):
    """在run内部指定位置拆分为两个run，返回 (前半run, 后半run)。

    split_pos: 在run文本中的字符位置。
    如果 split_pos <= 0 或 >= 文本长度，不拆分，返回 (None, run) 或 (run, None)。
    """
    run_text = ''.join(t.text or '' for t in run.iter(qn('w:t')))
    if split_pos <= 0:
        return None, run
    if split_pos >= len(run_text):
        return run, None

    # 复制 rPr
    rPr = run.find(qn('w:rPr'))

    # 前半 run
    before_run = OxmlElement('w:r')
    if rPr is not None:
        before_run.append(rPr.__copy__())
    before_t = OxmlElement('w:t')
    before_t.text = run_text[:split_pos]
    before_t.set(qn('xml:space'), 'preserve')
    before_run.append(before_t)

    # 后半 run
    after_run = OxmlElement('w:r')
    if rPr is not None:
        after_run.append(rPr.__copy__())
    after_t = OxmlElement('w:t')
    after_t.text = run_text[split_pos:]
    after_t.set(qn('xml:space'), 'preserve')
    after_run.append(after_t)

    # 替换原 run
    run_idx = list(target_elem).index(run)
    target_elem.remove(run)
    target_elem.insert(run_idx, before_run)
    target_elem.insert(run_idx + 1, after_run)

    return before_run, after_run


def _add_highlight_to_run(run, color='yellow'):
    """给指定run添加高亮底色（当前暂不启用，保留接口以备后续使用）。"""
    return


def _apply_comments_to_doc(doc, comment_list):
    """在文档中统一添加批注。

    comment_list: [(text_prefix, comment_text, anchor_type), ...]
    - text_prefix: 用于匹配段落的文本前缀
    - comment_text: 批注内容
    - anchor_type: 锚定方式，决定批注覆盖范围和高亮颜色
        'trailing_punct'  - 末尾标点错误，只标末尾标点字符，红色高亮
        'missing_end_punct' - 缺少句末标点，标段落末尾几个字，黄色高亮
        'number_prefix'   - 编号格式问题，只标编号部分，黄色高亮
        'heading_skip'    - 标题跳级，只标编号部分，黄色高亮
        'multi_level_num' - 多级编号，只标多级编号部分，黄色高亮
        'verb_after_num'  - 编号后跟动词，标编号+动词，黄色高亮
        'full_para'       - 整段问题，标整段，黄色高亮
    """
    if not comment_list:
        return

    body = doc.element.body

    # 收集所有段落及其文本
    para_map = []  # [(child_index, para_element, text), ...]
    child_idx = 0
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            text = ''.join(t.text for t in child.iter(qn('w:t')) if t.text)
            para_map.append((child_idx, child, text))
        child_idx += 1

    # 创建 comments XML
    comments_xml = (
        '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '</w:comments>'
    )
    comments_element = etree.fromstring(comments_xml.encode('utf-8'))
    next_id = 0
    now_str = datetime.datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')
    author = '格式化审查'

    # 已匹配的段落+批注组合集合
    matched_keys = set()

    for item in comment_list:
        if len(item) == 3:
            text_prefix, comment_text, anchor_type = item
        else:
            text_prefix, comment_text = item
            anchor_type = 'full_para'

        # 在输出段落中找到匹配的段落
        target_elem = None
        for ci, pelem, ptext in para_map:
            key = (ci, comment_text)
            if key in matched_keys:
                continue
            if ptext.startswith(text_prefix):
                target_elem = pelem
                matched_keys.add(key)
                break
        if target_elem is None:
            continue

        comment_id = str(next_id)

        # 1. 创建 w:comment 元素
        comment_elem = OxmlElement('w:comment')
        comment_elem.set(qn('w:id'), comment_id)
        comment_elem.set(qn('w:author'), author)
        comment_elem.set(qn('w:date'), now_str)
        comment_elem.set(qn('w:initials'), 'GS')

        p_comment = OxmlElement('w:p')
        r_comment = OxmlElement('w:r')
        t_comment = OxmlElement('w:t')
        t_comment.text = comment_text
        t_comment.set(qn('xml:space'), 'preserve')
        r_comment.append(t_comment)
        p_comment.append(r_comment)
        comment_elem.append(p_comment)
        comments_element.append(comment_elem)

        # 2. 根据 anchor_type 决定锚点位置和高亮范围
        runs = target_elem.findall(qn('w:r'))
        hl_color = 'red' if anchor_type == 'trailing_punct' else 'yellow'

        if anchor_type == 'trailing_punct' and runs:
            last_run = runs[-1]
            last_text = ''.join(t.text or '' for t in last_run.iter(qn('w:t')))
            if last_text and not last_text[-1].isalnum() and not last_text[-1].isspace():
                before, after = _split_run_at(target_elem, last_run, len(last_text) - 1)
                if after is not None:
                    _add_highlight_to_run(after, hl_color)
                    after_idx = list(target_elem).index(after)
                    cs = OxmlElement('w:commentRangeStart')
                    cs.set(qn('w:id'), comment_id)
                    target_elem.insert(after_idx, cs)
                    ce = OxmlElement('w:commentRangeEnd')
                    ce.set(qn('w:id'), comment_id)
                    target_elem.insert(after_idx + 2, ce)
                else:
                    _add_highlight_to_run(last_run, hl_color)
                    run_idx = list(target_elem).index(last_run)
                    cs = OxmlElement('w:commentRangeStart')
                    cs.set(qn('w:id'), comment_id)
                    target_elem.insert(run_idx, cs)
                    ce = OxmlElement('w:commentRangeEnd')
                    ce.set(qn('w:id'), comment_id)
                    target_elem.insert(run_idx + 2, ce)
            else:
                _add_highlight_to_run(last_run, hl_color)
                run_idx = list(target_elem).index(last_run)
                cs = OxmlElement('w:commentRangeStart')
                cs.set(qn('w:id'), comment_id)
                target_elem.insert(run_idx, cs)
                ce = OxmlElement('w:commentRangeEnd')
                ce.set(qn('w:id'), comment_id)
                target_elem.insert(run_idx + 2, ce)

        elif anchor_type == 'missing_end_punct' and runs:
            last_run = runs[-1]
            last_text = ''.join(t.text or '' for t in last_run.iter(qn('w:t')))
            highlight_len = min(3, len(last_text))
            if highlight_len > 0 and len(last_text) > highlight_len:
                before, after = _split_run_at(target_elem, last_run, len(last_text) - highlight_len)
                if after is not None:
                    _add_highlight_to_run(after, hl_color)
                    after_idx = list(target_elem).index(after)
                    cs = OxmlElement('w:commentRangeStart')
                    cs.set(qn('w:id'), comment_id)
                    target_elem.insert(after_idx, cs)
                    ce = OxmlElement('w:commentRangeEnd')
                    ce.set(qn('w:id'), comment_id)
                    target_elem.insert(after_idx + 2, ce)
                else:
                    _add_highlight_to_run(last_run, hl_color)
                    run_idx = list(target_elem).index(last_run)
                    cs = OxmlElement('w:commentRangeStart')
                    cs.set(qn('w:id'), comment_id)
                    target_elem.insert(run_idx, cs)
                    ce = OxmlElement('w:commentRangeEnd')
                    ce.set(qn('w:id'), comment_id)
                    target_elem.insert(run_idx + 2, ce)
            else:
                _add_highlight_to_run(last_run, hl_color)
                run_idx = list(target_elem).index(last_run)
                cs = OxmlElement('w:commentRangeStart')
                cs.set(qn('w:id'), comment_id)
                target_elem.insert(run_idx, cs)
                ce = OxmlElement('w:commentRangeEnd')
                ce.set(qn('w:id'), comment_id)
                target_elem.insert(run_idx + 2, ce)

        elif anchor_type in ('number_prefix', 'heading_skip') and runs:
            first_run = runs[0]
            first_text = ''.join(t.text or '' for t in first_run.iter(qn('w:t')))
            num_match = re.match(
                r'^(\d+[.、．]\s*|[一二三四五六七八九十]+[、]\s*|（[一二三四五六七八九十]+）\s*|（\d+）\s*)',
                first_text
            )
            if num_match and len(num_match.group(0)) < len(first_text):
                before, after = _split_run_at(target_elem, first_run, len(num_match.group(0)))
                if before is not None:
                    _add_highlight_to_run(before, hl_color)
                    before_idx = list(target_elem).index(before)
                    cs = OxmlElement('w:commentRangeStart')
                    cs.set(qn('w:id'), comment_id)
                    target_elem.insert(before_idx, cs)
                    ce = OxmlElement('w:commentRangeEnd')
                    ce.set(qn('w:id'), comment_id)
                    target_elem.insert(before_idx + 2, ce)
                else:
                    _add_highlight_to_run(first_run, hl_color)
                    run_idx = list(target_elem).index(first_run)
                    cs = OxmlElement('w:commentRangeStart')
                    cs.set(qn('w:id'), comment_id)
                    target_elem.insert(run_idx, cs)
                    ce = OxmlElement('w:commentRangeEnd')
                    ce.set(qn('w:id'), comment_id)
                    target_elem.insert(run_idx + 2, ce)
            elif num_match:
                _add_highlight_to_run(first_run, hl_color)
                run_idx = list(target_elem).index(first_run)
                cs = OxmlElement('w:commentRangeStart')
                cs.set(qn('w:id'), comment_id)
                target_elem.insert(run_idx, cs)
                ce = OxmlElement('w:commentRangeEnd')
                ce.set(qn('w:id'), comment_id)
                target_elem.insert(run_idx + 2, ce)
            else:
                _add_highlight_to_run(first_run, hl_color)
                run_idx = list(target_elem).index(first_run)
                cs = OxmlElement('w:commentRangeStart')
                cs.set(qn('w:id'), comment_id)
                target_elem.insert(run_idx, cs)
                ce = OxmlElement('w:commentRangeEnd')
                ce.set(qn('w:id'), comment_id)
                target_elem.insert(run_idx + 2, ce)

        elif anchor_type == 'multi_level_num' and runs:
            first_run = runs[0]
            first_text = ''.join(t.text or '' for t in first_run.iter(qn('w:t')))
            multi_match = re.match(r'^(\d+[.、．]\d+[.、．]?\s*)', first_text)
            if multi_match and len(multi_match.group(0)) < len(first_text):
                before, after = _split_run_at(target_elem, first_run, len(multi_match.group(0)))
                if before is not None:
                    _add_highlight_to_run(before, hl_color)
                    before_idx = list(target_elem).index(before)
                    cs = OxmlElement('w:commentRangeStart')
                    cs.set(qn('w:id'), comment_id)
                    target_elem.insert(before_idx, cs)
                    ce = OxmlElement('w:commentRangeEnd')
                    ce.set(qn('w:id'), comment_id)
                    target_elem.insert(before_idx + 2, ce)
                else:
                    _add_highlight_to_run(first_run, hl_color)
                    run_idx = list(target_elem).index(first_run)
                    cs = OxmlElement('w:commentRangeStart')
                    cs.set(qn('w:id'), comment_id)
                    target_elem.insert(run_idx, cs)
                    ce = OxmlElement('w:commentRangeEnd')
                    ce.set(qn('w:id'), comment_id)
                    target_elem.insert(run_idx + 2, ce)
            elif multi_match:
                _add_highlight_to_run(first_run, hl_color)
                run_idx = list(target_elem).index(first_run)
                cs = OxmlElement('w:commentRangeStart')
                cs.set(qn('w:id'), comment_id)
                target_elem.insert(run_idx, cs)
                ce = OxmlElement('w:commentRangeEnd')
                ce.set(qn('w:id'), comment_id)
                target_elem.insert(run_idx + 2, ce)
            else:
                _add_highlight_to_run(first_run, hl_color)
                run_idx = list(target_elem).index(first_run)
                cs = OxmlElement('w:commentRangeStart')
                cs.set(qn('w:id'), comment_id)
                target_elem.insert(run_idx, cs)
                ce = OxmlElement('w:commentRangeEnd')
                ce.set(qn('w:id'), comment_id)
                target_elem.insert(run_idx + 2, ce)

        elif anchor_type == 'verb_after_num' and runs:
            first_run = runs[0]
            first_text = ''.join(t.text or '' for t in first_run.iter(qn('w:t')))
            verb_match = re.match(
                r'^(\d+[.、．]\s*[\u662f\u8981\u6709\u80fd\u5e94\u4f1a\u53ef\u5c06\u5f97\u5fc5\u987b]{1,2})',
                first_text
            )
            if verb_match and len(verb_match.group(0)) < len(first_text):
                before, after = _split_run_at(target_elem, first_run, len(verb_match.group(0)))
                if before is not None:
                    _add_highlight_to_run(before, hl_color)
                    before_idx = list(target_elem).index(before)
                    cs = OxmlElement('w:commentRangeStart')
                    cs.set(qn('w:id'), comment_id)
                    target_elem.insert(before_idx, cs)
                    ce = OxmlElement('w:commentRangeEnd')
                    ce.set(qn('w:id'), comment_id)
                    target_elem.insert(before_idx + 2, ce)
                else:
                    _add_highlight_to_run(first_run, hl_color)
                    run_idx = list(target_elem).index(first_run)
                    cs = OxmlElement('w:commentRangeStart')
                    cs.set(qn('w:id'), comment_id)
                    target_elem.insert(run_idx, cs)
                    ce = OxmlElement('w:commentRangeEnd')
                    ce.set(qn('w:id'), comment_id)
                    target_elem.insert(run_idx + 2, ce)
            else:
                num_match = re.match(r'^(\d+[.、．]\s*)', first_text)
                end_pos = len(num_match.group(0)) if num_match else min(len(first_text), 5)
                if end_pos < len(first_text):
                    before, after = _split_run_at(target_elem, first_run, end_pos)
                    if before is not None:
                        _add_highlight_to_run(before, hl_color)
                        before_idx = list(target_elem).index(before)
                        cs = OxmlElement('w:commentRangeStart')
                        cs.set(qn('w:id'), comment_id)
                        target_elem.insert(before_idx, cs)
                        ce = OxmlElement('w:commentRangeEnd')
                        ce.set(qn('w:id'), comment_id)
                        target_elem.insert(before_idx + 2, ce)
                    else:
                        _add_highlight_to_run(first_run, hl_color)
                        run_idx = list(target_elem).index(first_run)
                        cs = OxmlElement('w:commentRangeStart')
                        cs.set(qn('w:id'), comment_id)
                        target_elem.insert(run_idx, cs)
                        ce = OxmlElement('w:commentRangeEnd')
                        ce.set(qn('w:id'), comment_id)
                        target_elem.insert(run_idx + 2, ce)
                else:
                    _add_highlight_to_run(first_run, hl_color)
                    run_idx = list(target_elem).index(first_run)
                    cs = OxmlElement('w:commentRangeStart')
                    cs.set(qn('w:id'), comment_id)
                    target_elem.insert(run_idx, cs)
                    ce = OxmlElement('w:commentRangeEnd')
                    ce.set(qn('w:id'), comment_id)
                    target_elem.insert(run_idx + 2, ce)

        else:
            # full_para 或其他：高亮整段
            for r in runs:
                _add_highlight_to_run(r, hl_color)
            pPr = target_elem.find(qn('w:pPr'))
            if pPr is not None:
                insert_idx = list(target_elem).index(pPr) + 1
            else:
                insert_idx = 0
            cs = OxmlElement('w:commentRangeStart')
            cs.set(qn('w:id'), comment_id)
            target_elem.insert(insert_idx, cs)
            ce = OxmlElement('w:commentRangeEnd')
            ce.set(qn('w:id'), comment_id)
            target_elem.append(ce)

        # 插入 commentReference run
        ref_run = OxmlElement('w:r')
        ref_rPr = OxmlElement('w:rPr')
        ref_rStyle = OxmlElement('w:rStyle')
        ref_rStyle.set(qn('w:val'), 'CommentReference')
        ref_rPr.append(ref_rStyle)
        ref_run.append(ref_rPr)
        ref_cr = OxmlElement('w:commentReference')
        ref_cr.set(qn('w:id'), comment_id)
        ref_run.append(ref_cr)
        ce_elem = target_elem.find(
            f'.//w:commentRangeEnd[@w:id="{comment_id}"]',
            namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        )
        if ce_elem is not None:
            ce_idx = list(target_elem).index(ce_elem)
            target_elem.insert(ce_idx + 1, ref_run)
        else:
            target_elem.append(ref_run)

        next_id += 1

    # 将 comments 保存为 Part
    comments_bytes = etree.tostring(comments_element, xml_declaration=True, encoding='UTF-8', standalone=True)
    doc_part = doc.part

    for rel in doc_part.rels.values():
        if 'comments' in rel.reltype:
            comments_part = rel.target_part
            comments_part._blob = comments_bytes
            return

    # 创建新的 Part
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    comments_partname = PackURI('/word/comments.xml')
    comments_part = Part(
        partname=comments_partname,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml',
        blob=comments_bytes,
        package=doc_part.package
    )
    doc_part.relate_to(comments_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments')


def _add_page_number(doc):
    """在文档页脚添加居中页码（— 1 —风格）。"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = fp._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')

    def make_run(text_or_field, is_field=False):
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), FONT_FANGSONG)
        rFonts.set(qn('w:eastAsia'), '宋体')
        rFonts.set(qn('w:hAnsi'), FONT_FANGSONG)
        rPr.append(rFonts)
        for tag in ('w:sz', 'w:szCs'):
            sz = OxmlElement(tag)
            sz.set(qn('w:val'), '28')
            rPr.append(sz)
        r.append(rPr)
        if is_field:
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), text_or_field)
            r.append(fld)
        else:
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = text_or_field
            r.append(t)
        return r

    def make_instrText(instr):
        r = OxmlElement('w:r')
        instrT = OxmlElement('w:instrText')
        instrT.set(qn('xml:space'), 'preserve')
        instrT.text = instr
        r.append(instrT)
        return r

    p_elem = fp._p
    p_elem.append(make_run('— ', False))
    p_elem.append(make_run('begin', True))
    p_elem.append(make_instrText(' PAGE '))
    p_elem.append(make_run('separate', True))
    p_elem.append(make_run('end', True))
    p_elem.append(make_run(' —', False))
